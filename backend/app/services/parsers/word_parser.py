"""Word document parser using python-docx."""

import io
from typing import List

from docx import Document
from docx.table import Table

from app.services.parsers.base import BaseParser, ParsedContent


class WordParser(BaseParser):
    """Parser for Word documents."""

    SUPPORTED_TYPES = ['docx', 'doc']

    def supports(self, file_type: str) -> bool:
        """Check if this parser supports the given file type."""
        return file_type.lower() in self.SUPPORTED_TYPES

    async def parse(self, file_path: str) -> ParsedContent:
        """Parse a Word document.
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            ParsedContent with extracted text, tables, and images
        """
        result = ParsedContent()

        try:
            doc = Document(file_path)

            # Extract document properties
            result.metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            }

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    result.text_blocks.append(text)

            # Extract tables
            for table in doc.tables:
                table_data = self._extract_table(table)
                if table_data:
                    result.tables.append(table_data)

            # Extract images
            images = self._extract_images(doc)
            result.images.extend(images)

        except Exception as e:
            result.metadata["error"] = str(e)

        return result

    def _extract_table(self, table: Table) -> List[List[str]]:
        """Extract data from a Word table."""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        return table_data

    def _extract_images(self, doc: Document) -> List[bytes]:
        """Extract images from a Word document."""
        images = []
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    images.append(image_part.blob)
        except Exception:
            pass
        return images
